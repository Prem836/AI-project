import os
import pypdf
import docx

def load_document(file_path: str):
    """
    Loads a PDF or DOCX document and extracts its text contents.
    Returns:
        list of dict: [{"text": str, "metadata": dict}]
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".docx":
        return _load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX files are supported.")

def _load_pdf(file_path: str):
    try:
        reader = pypdf.PdfReader(file_path)
        total_pages = len(reader.pages)
        full_text = []
        pages_data = []
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            full_text.append(page_text)
            pages_data.append({
                "page_num": i + 1,
                "text": page_text
            })
            
        concatenated_text = "\n".join(full_text)
        return [{
            "text": concatenated_text,
            "metadata": {
                "source": file_path,
                "total_pages": total_pages,
                "pages": pages_data
            }
        }]
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF file {file_path}: {str(e)}")

def _load_docx(file_path: str):
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        concatenated_text = "\n".join(full_text)
        return [{
            "text": concatenated_text,
            "metadata": {
                "source": file_path,
                "total_pages": 1,
                "pages": [
                    {
                        "page_num": 1,
                        "text": concatenated_text
                    }
                ]
            }
        }]
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX file {file_path}: {str(e)}")

